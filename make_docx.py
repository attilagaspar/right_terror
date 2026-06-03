#!/usr/bin/env python3
"""Convert main_rr.tex to main_rr.docx for typesetting."""

import re, zipfile, io, copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\agaspar\Dropbox\research\rightterror\draft\main_rr.docx"

# ── Citation map ──────────────────────────────────────────────
CITE = {
    'walterschang': 'Walters and Chang 2021',
    'auger2020right': 'Auger 2020',
    'berrebi2008voters': 'Berrebi and Klor 2008',
    'legewie2013terrorist': 'Legewie 2013',
    'montalvo2011voting': 'Montalvo 2011',
    'godefroidt2023terrorism': 'Godefroidt 2023',
    'jakobsson2014did': 'Jakobsson and Blom 2014',
    'pickard2021s': 'Pickard et al. 2022',
    'marques1988': 'Marques, Yzerbyt, and Leyens 1988',
    'krause2023does': 'Krause and Matsunaga 2023',
    'sabet2025terrorism': 'Sabet, Liebald, and Friebel 2025',
    'volker2024terrorist': 'Völker 2024',
    'cossu2025terror': 'Cossu and Froio 2025',
    'jacobs2021not': 'Jacobs and van Spanje 2021',
    'miklosi2011': 'Miklósi 2011',
    'jaszberenyi': 'Jászberényi 2016',
    'tamas': 'Tamás 2019',
    'gyongyosi2020financial': 'Gyöngyösi and Verner 2020',
    'karacsony2011secret': 'Karácsony and Róna 2011',
    'mti2010': 'MTI 2010',
    'mti2011': 'MTI 2011',
    'national_election_commission_data': 'National Election Commission 2024',
    'tstar_database': 'Central Statistical Office 2018',
    'census_2001': 'Central Statistical Office 2001',
    'cinelli2020making': 'Cinelli and Hazlett 2020',
    'agerberg2021personal': 'Agerberg and Sohlberg 2021',
    'bohmelt2020can': 'Böhmelt, Bove, and Nussio 2020',
    'simonovits2018seeing': 'Simonovits, Kezdi, and Kardos 2018',
    'szombati2018revolt': 'Szombati 2018',
    'hager2019ethnic': 'Hager, Krakowski, and Schaub 2019',
    'shoshani2008drama': 'Shoshani and Slone 2008',
    'jacobs2023': 'Jacobs and van Spanje 2023',
    'huff2018public': 'Huff and Kertzer 2018',
    'Amnesty_International_2010_EUR710072010en': 'Amnesty International 2010',
    'BBC_2011_bulgarian_rally_roma': 'BBC News 2011',
}

REFS = {
    'app:context': 'ONLINE APPENDIX SECTION A',
    'app:farright': 'ONLINE APPENDIX SECTION B',
    'app:robustnes': 'ONLINE APPENDIX SECTION G',
    'tab:locations': 'TABLE OA1',
    'table_balance': 'TABLE OA4',
    'fig_map': 'FIGURE OA4',
    'table1': 'TABLE 1',
    'table1_app': 'TABLE OA5',
    'fig_eventstudy': 'FIGURE 1',
    'table_event': 'TABLE OA7',
    'fig2': 'FIGURE 2',
    'eq1': 'EQUATION 1',
}

def citep_repl(m):
    keys = [k.strip() for k in m.group(1).split(',')]
    return '(' + '; '.join(CITE.get(k, k) for k in keys) + ')'

def ref_repl(m):
    return REFS.get(m.group(1).strip(), m.group(1).upper())

def clean_math(t):
    subs = [
        (r'\\text\{([^}]*)\}', r'\1'),
        (r'\\mathbf\{([^}]*)\}', r'\1'),
        (r'\\varepsilon', 'ε'), (r'\\gamma', 'γ'), (r'\\beta', 'β'),
        (r'\\times', '×'), (r'\\hat\{([^}]*)\}', r'\1̂'),
        (r'\{', ''), (r'\}', ''), (r'\\_', '_'), (r'_\{it\}', 'ᵢₜ'),
        (r'_\{i\}', 'ᵢ'), (r'_\{t\}', 'ₜ'), (r'_0', '₀'),
        (r'_1', '₁'), (r'_2', '₂'), (r'_3', '₃'),
    ]
    for pat, rep in subs:
        t = re.sub(pat, rep, t)
    return t

def clean(text):
    # Citations with optional page note
    text = re.sub(
        r'\\citep\[([^\]]*)\]\{([^}]+)\}',
        lambda m: '(' + '; '.join(CITE.get(k.strip(), k.strip()) for k in m.group(2).split(','))
                  + ', ' + m.group(1) + ')',
        text)
    text = re.sub(r'\\citep\{([^}]+)\}', citep_repl, text)
    text = re.sub(r'\\citeauthor\{([^}]+)\}',
                  lambda m: CITE.get(m.group(1).strip(), m.group(1)).rsplit(' ', 1)[0], text)
    text = re.sub(r'\\citeyear\{([^}]+)\}',
                  lambda m: CITE.get(m.group(1).strip(), m.group(1)).rsplit(' ', 1)[-1], text)
    text = re.sub(r'\\citeapos\{([^}]+)\}',
                  lambda m: CITE.get(m.group(1).strip(), m.group(1)).rsplit(' ', 1)[0]
                            + "’s ("
                            + CITE.get(m.group(1).strip(), m.group(1)).rsplit(' ', 1)[-1] + ')',
                  text)
    text = re.sub(r'\\ref\{([^}]+)\}', ref_repl, text)
    # Inline math
    text = re.sub(r'\$([^$]+)\$', lambda m: clean_math(m.group(1)), text)
    # Formatting commands
    for cmd in ('textit', 'textbf', 'emph', 'textsc', 'textrm', 'textnormal',
                'small', 'footnotesize', 'normalsize', 'noindent'):
        text = re.sub(r'\\' + cmd + r'\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\' + cmd + r'\b\s*', '', text)
    # Special chars
    text = text.replace('\\%', '%').replace('\\$', '$').replace('\\&', '&')
    text = text.replace('---', '—').replace('--', '–')
    text = text.replace('``', '“').replace("''", '”')
    text = text.replace('\\pm', '±')
    text = re.sub(r'\\url\{[^}]*\}', '', text)
    # Strip remaining LaTeX commands
    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z@]+\*?\s*', '', text)
    text = re.sub(r'\{([^}]*)\}', r'\1', text)
    text = re.sub(r'[ \t]+', ' ', text).strip()
    return text

# ── Footnote builder ──────────────────────────────────────────
# We collect footnotes and inject into the docx zip at the end.

footnotes_data = []   # list of cleaned footnote strings

def add_fn(para, text):
    """Append footnote marker to paragraph and register footnote text."""
    fn_id = len(footnotes_data) + 1  # footnotes start at id=1
    footnotes_data.append(text)
    r = para.add_run()
    fref = OxmlElement('w:footnoteReference')
    fref.set(qn('w:id'), str(fn_id))
    rpr = OxmlElement('w:rPr')
    style_el = OxmlElement('w:rStyle')
    style_el.set(qn('w:val'), 'FootnoteReference')
    rpr.append(style_el)
    r._r.insert(0, rpr)
    r._r.append(fref)

def add_fn_mid(para, before_text, fn_text, after_text=''):
    """Add text + footnote marker + optional trailing text to paragraph."""
    if before_text:
        para.add_run(before_text)
    add_fn(para, fn_text)
    if after_text:
        para.add_run(after_text)

# ── Document builder ──────────────────────────────────────────

doc = Document()

# --- Title ---
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH = 1  # center
run = t.add_run('Right-wing terrorism and far-right support: Evidence from anti-Roma attacks in Hungary')
run.bold = True
run.font.size = Pt(14)

# --- Authors with affiliation footnotes ---
a = doc.add_paragraph()
a.alignment = 1
a.add_run('Gábor Simonovits')
add_fn(a, 'Associate Professor, Department of Political Science, Central European University and Senior Research Fellow at the ELTE Centre for Social Sciences, Institute for Political Science.')
a.add_run(', Attila Gáspár')
add_fn(a, 'Research Fellow at the ELTE Centre for Economic and Regional Studies (ELTE CERS).')
a.add_run(', Gábor Békés')
add_fn(a, 'Associate Professor, Department of Economics, Central European University, Senior Research Fellow at the ELTE Centre for Economic and Regional Studies (ELTE CERS) and Research fellow, CEPR.')
a.add_run(', Márton Végh')
add_fn(a, 'Independent researcher.')

# --- Abstract ---
doc.add_paragraph()
abst_label = doc.add_paragraph()
abst_label.add_run('Abstract').bold = True

abst = doc.add_paragraph()
abst.add_run(
    'How do ethnically motivated terrorist attacks shape electoral support for the far right? '
    'We study a unique case: a coordinated series of anti-Roma murders in Hungary in 2008–2009, '
    'the most severe episode of anti-minority violence in the country since World War II. '
    'Combining difference-in-differences and synthetic control methods, we compare attacked settlements '
    'to multiple counterfactuals, including planned-but-unrealized targets. We find that Jobbik, '
    'Hungary’s radical right party, gained 11–14 percentage points more support in attacked '
    'villages than in comparable controls in the 2010 election—an increase 53–70% larger than '
    'baseline trends. The effect persisted for several years and spilled over to nearby settlements. '
    'In contrast to some research from Western Europe suggesting that right-wing terrorism can reduce '
    'far-right appeal, our findings highlight how deep-seated prejudice can reverse this pattern. '
    'The results underscore the importance of antecedent inter-group relations in conditioning '
    'political reactions to ethnic violence.'
)
add_fn(abst,
       'The authors declare no competing interests. The views expressed in this article are those '
       'of the authors and do not necessarily reflect the views of their employers and affiliated institutions.')

# --- Section: Introduction ---
doc.add_page_break()
h1 = doc.add_paragraph()
h1.add_run('Introduction').bold = True
h1.style = doc.styles['Heading 1']

p = doc.add_paragraph()
p.add_run(
    'The last two decades have seen an unprecedented rise in violence committed by right-wing extremists. '
    'This tendency is demonstrated by high profile events such as the 2011 attacks in Norway, the 2019 '
    'shootings in Halle and Christchurch as well as systematically collected data showing the spread of '
    'such terror-attacks in the Western world (Walters and Chang 2021; Auger 2020). A key question '
    'regarding these acts of violence is their relationship to the growing popularity of far-right political '
    'parties: to what extent does violence against vulnerable minorities fuel support for far-right parties?'
)

p2 = doc.add_paragraph()
p2.add_run(
    'While there is a rich literature exploring the impact of terrorism on public opinion '
    '(Berrebi and Klor 2008; Legewie 2013; Montalvo 2011) most studies consider terror attacks '
    'committed by minority groups (Godefroidt 2023). In contrast, the literature on the impact '
    'of right-wing terror attacks is much sparser, leading to mixed findings. On the one hand, some '
    'studies suggest that right-wing terror attacks increase affect for immigrants and lower approval '
    'of right-wing positions in their aftermath (Jakobsson and Blom 2014; Pickard et al. 2022). This '
    'is usually attributed to the “Black Sheep effect”, whereby ingroup members reject deviants '
    'to maintain group status (Marques, Yzerbyt, and Leyens 1988) and thus voters reject parties and '
    'ideas that they associate with the perpetrators who challenge existing social norms. In contrast, '
    'other studies have found the opposite, showing that at least under some circumstances such '
    'terror-attacks increase support for the far-right, possibly because they raise the salience of '
    'issues – such as immigration or law-and-order – owned by these parties '
    '(Krause and Matsunaga 2023; Sabet, Liebald, and Friebel 2025; Völker 2024).'
)
add_fn(p2,
       'See also (Cossu and Froio 2025) for a recent analysis of how terror attacks have shaped '
       'public discourse in France.')

p3 = doc.add_paragraph()
p3.add_run(
    'However, it is important to note that the scope of cases covered by existing studies are quite '
    'narrow. They focus almost entirely on Western democracies – like Germany and Norway – '
    'and explore the consequences of terror-attacks against immigrants. Because the impact of such '
    'violent events clearly depends on both the characteristics of the perpetrators and victims '
    '(Jacobs and van Spanje 2021) and the way the media covers such events '
    '(Sabet, Liebald, and Friebel 2025; Völker 2024), expanding the universe of cases is essential '
    'to a more complete picture. The contribution of this research note is to bring new evidence to '
    'this question by studying a new case and utilizing a novel methodological approach that can '
    'provide more credible causal estimates.'
)

p4 = doc.add_paragraph()
p4.add_run(
    'Specifically, we explore the electoral impact of a series of ethnically motivated terrorist '
    'attacks against members of the Roma ethnic community by right-wing extremists in Hungary in '
    '2008-2009. The attacks strongly resembled terrorist attacks in other European countries and took '
    'place against the backdrop of growing ethnic tensions and the rise of the anti-Roma radical right '
    'Jobbik party that increased its support from 2.2% in 2006 to over 16% by 2010. Utilizing an '
    'ensemble of approaches to causal inference that combine both quasi-experimental and synthetic '
    'control groups with longitudinal data on the settlement-level vote share of Jobbik, we '
    'demonstrate that support for the far-right increased by 11 to 14 percentage points more in the '
    'attacked settlements from 2006 to 2010 than it did in the control settlements. Our analysis '
    'also provides evidence of the temporal and spatial scope of this effect: the effect of the '
    'attacks spilled over to neighboring villages, even quite distant from the location of the '
    'attacks, and persisted till the national and EU elections in 2014, but not in the national '
    'election of 2018.'
)

# --- Section: Research Design ---
h_rd = doc.add_paragraph('Research design')
h_rd.style = doc.styles['Heading 1']

# -- Subsection: Context --
h_ctx = doc.add_paragraph('Context')
h_ctx.style = doc.styles['Heading 2']

p_ctx1 = doc.add_paragraph()
p_ctx1.add_run('We study the impact of anti-Roma terrorism on far-right political support in Hungary.')
add_fn(p_ctx1,
       'In ONLINE APPENDIX SECTION A we show survey evidence that they face serious prejudice '
       'in Hungary as well as in the rest of Europe and have been increasingly facing ethnic violence '
       'since the Great Recession.')
p_ctx1.add_run(
    ' The anti-Roma attacks constituted the single most important case of anti-minority violence in '
    'the country since World War II. They took place in distant villages across Eastern and Northern '
    'Hungary over the course of approximately one year from July 2008 to August 2009, leaving six '
    'people dead and many more injured. The terrorists were a group of men who consistently employed '
    'the same methods and selected similar targets: they attacked Roma-inhabited houses with arson '
    'and gunfire at multiple locations. Details are found in TABLE OA1. The trial proceedings '
    'revealed that the murders were carefully premeditated against civilians with whom the attackers '
    'had no previous relations. The motives of the perpetrators were clearly political in that they '
    'sought to incite further violence among and against Roma (Miklósi 2011; Jászberényi 2016; Tamás 2019).'
)

p_ctx2 = doc.add_paragraph()
p_ctx2.add_run(
    'Meanwhile, the political landscape also went through radical realignment. The immediate aftermath '
    'of the financial crisis and the ensuing collapse of the governing socialist-liberal coalition '
    'government brought not only simmering ethnic tensions but also the rise of the organized radical '
    'right (Gyöngyösi and Verner 2020). The radical Jobbik party (“Better” or “More '
    'Right” in Hungarian) grew from a marginal organization in 2006 into a sizable parliamentary '
    'force, capturing 15% of the national vote in the 2009 European Parliament elections just before '
    'the last attacks, and 17% in the 2010 parliamentary elections. Anti-Gypsism was at the center '
    'of Jobbik’s political agenda, as the promise of putting an end to “Gypsy crime” '
    'and Roma welfare scrounging featured prominently in the party’s platform. It allowed Jobbik '
    'to successfully mobilize voters in Hungary’s Northern industrial rustbelt and left-behind '
    'rural areas on the Northern Great Plains, the regions with the highest Roma population share '
    '(Karácsony and Róna 2011). We provide additional details about anti-Roma prejudice in Hungary '
    'as well as the rise of Jobbik compared to general EU wide trends in the Online Appendix Sections '
    'ONLINE APPENDIX SECTION A and ONLINE APPENDIX SECTION B.'
)

# -- Subsection: Identification --
h_id = doc.add_paragraph('Identification')
h_id.style = doc.styles['Heading 2']

p_id1 = doc.add_paragraph()
p_id1.add_run(
    'To estimate the causal effect of the attacks on electoral results we compare changes in the '
    'support for Jobbik in attacked villages to similar but unattacked ones. We do so relying on '
    'versions of the following difference-in-differences model:'
)

# Equation block
eq = doc.add_paragraph()
eq.alignment = 1  # center
eq.add_run(
    'EQUATION 1:  JobbikShareᵢₜ = β₀ + β₁ Attackedᵢ + β₂ Postₜ + β₃ (Attackedᵢ × Postₜ) + Xᵢₜγ + εᵢₜ'
).italic = True

p_id2 = doc.add_paragraph()
p_id2.add_run(
    'where JobbikShareᵢₜ is the vote share of Jobbik in settlement i at time t, '
    'Attackedᵢ is a dummy variable indicating whether settlement i was attacked, '
    'Postₜ is a dummy variable equal to 1 if t=2010 (post-treatment period) and 0 if t=2006 '
    '(pre-treatment period), '
    'Attackedᵢ × Postₜ is the difference-in-differences interaction term to estimate the average '
    'treatment effect on the treated (ATET) and εᵢₜ is the error term.'
)

p_id3 = doc.add_paragraph()
p_id3.add_run(
    'The causal effect of the attacks on subsequent election outcomes is identified by the coefficient '
    'β₃ if the parallel trend assumption is met, that is, if Jobbik support in attacked villages '
    'would have evolved over a similar path compared to the rest of the country in the absence of '
    'the attacks. A key challenge to this assumption is selection: terrorists chose locations with '
    'rising frictions between the Roma and the non-Roma population, we expect counterfactual Jobbik '
    'support to be also higher in the attacked group than in the control group. This would cause an '
    'upward bias in our estimate β̂₃.'
)

p_id4 = doc.add_paragraph()
p_id4.add_run(
    'To address this challenge, we rely on different control groups that could plausibly be used to '
    'estimate counterfactual changes in attacked villages in the absence of the attacks. Specifically, '
    'we first rely on a '
)
run_it = p_id4.add_run('design-based')
run_it.italic = True
p_id4.add_run(
    ' control group which includes settlements where the terrorists were later – during the '
    'court phase – revealed to have planned further attacks, but had no opportunity to carry '
    'them out due to their capture.'
)
add_fn(p_id4,
       'After the terrorists were identified and captured, the police uncovered evidence that they '
       'were planning to attack Roma at additional localities across the same geographic area. Police '
       'press releases on the planned attacks have widely been reported in the Hungarian media, '
       'including the official Hungarian News Agency Corporation (MTI 2010; MTI 2011). More '
       'information on planned attacks is found in TABLE OA1.')
p_id4.add_run(
    ' The advantage of this comparison is that they were arguably selected on the same unobservables '
    'as the attacked variables. The disadvantage is that there are only four such settlements.'
)

p_id5 = doc.add_paragraph()
p_id5.add_run(
    'Second, we also we create a synthetic control equivalent for each of the 9 towns and villages, '
    'and the control group of the regression are the synthetically generated control settlements. '
    'We use the set of similar settlements as the '
)
p_id5.add_run('donor pool').italic = True
p_id5.add_run(
    '. The synthetic control settlements arguably have the socio-economic fundamentals closest to '
    'the attacked settlements. A further advantage of the syntethic control approach is that it also '
    'allows us to assess the impact of '
)
p_id5.add_run('each').italic = True
p_id5.add_run(
    ' terror attach separately. We use these settlement level estimates to provide further evidence '
    'of the validity of our assumptions.'
)

p_id6 = doc.add_paragraph()
p_id6.add_run(
    'Finally, we also present baseline estimates from a '
)
p_id6.add_run('broad control group').italic = True
p_id6.add_run(
    ' that includes every settlement in Hungary and a '
)
p_id6.add_run('restricted control group').italic = True
p_id6.add_run(
    ' that includes settlements that had at most 10.000 inhabitants at the time and had a Roma '
    'population (the restricted control group also serves as the donor pool for the synthetic '
    'control analysis). In these specifications we also control for a vector of covariates to '
    'further mitigate selection. This vector includes the share of the Roma population according '
    'to the most recent (2001) census, and indicators of economic distress: the settlement level '
    'unemployment rate, the number of crimes reported, and the number of inhabitants with a '
    'criminal record.'
)

# -- Subsection: Data --
h_data = doc.add_paragraph('Data')
h_data.style = doc.styles['Heading 2']

p_data = doc.add_paragraph()
p_data.add_run(
    'We rely on data from several sources. Election data come from the National Election '
    'Commission’s website (National Election Commission 2024). Settlement-level variables '
    'come from the TEIR system collected by the Hungarian Central Statistical Office, Hungarian '
    'State Treasury, National Tax and Customs Administration and the Ministry for Home Affairs and '
    'their legal successors (Central Statistical Office 2018).'
)
add_fn(p_data, 'The data provided by different data owners were harmonized by the ELTE KRTK Databank.')
p_data.add_run(
    ' The number of citizens identifying as Roma comes from the 2001 Census (Central Statistical '
    'Office 2001).'
)
add_fn(p_data,
       'The Census data may miss recent Roma shifts, confounding attacks with far-right support. '
       'To address this issue use contemporaneous settlement data and add proxies of change, with '
       'new estimates if anything, implying downward bias in the baseline.')
p_data.add_run(
    ' Finally, we also include the number of times the settlement was mentioned on the far-right '
    'news portal kuruc.info’s “gypsy crime” column in 2008 before the first attack '
    'took place. We scraped all posts from the far-right '
)
p_data.add_run('kuruc.info').italic = True
p_data.add_run(
    ' news site’s “gypsy crime” column between 2006 and 2018 to gain a '
    'settlement-level contemporaneous indicator of tensions between the Roma and the non-Roma.'
)
add_fn(p_data,
       'This Hungarian far-right online outlet is registered in the US; its ownership structure '
       'and authors are officially unknown due to First Amendment protections on free speech It '
       'has been widely alleged that the portal had close ties to and was edited by senior Jobbik '
       'members. The Kuruc.info “cigánybűnözés” (“Gypsy '
       'crime”) section is racist propaganda tool that selectively reports crimes allegedly '
       'committed by Roma individuals to reinforce negative stereotypes, fueling prejudice and '
       'justifying discrimination against the Roma community. We looked for mentions of Hungarian '
       'settlement names in all posts between 2006 and 2018, and used the number of mentions as '
       'a proxy for ethnic tensions.')

p_data2 = doc.add_paragraph()
p_data2.add_run(
    'TABLE OA4 in the Appendix shows the demographic and economic profile of attacked settlements '
    'and different sets of control settlements. There we list the 2006 values of all variables that '
    'were used to construct the synthetic controls. We show the mean of the variables for the '
    'attacked settlements (column 1), the rest of the country (column 3), the donor pool '
    '(settlements with at most 10.000 inhabitants, of which at least 1 identified as Roma in the '
    '2001 census, column 5), and the average of those settlements that were eventually used for '
    'generating the synthetic control settlements. FIGURE OA4 in the Appendix shows the '
    'geographical location of attached, planned attack villages and similar settlements. These '
    'units will also serve as the donor pool for our synthetic control method.'
)

# --- Section: Results ---
h_res = doc.add_paragraph('Results')
h_res.style = doc.styles['Heading 1']

p_res1 = doc.add_paragraph()
p_res1.add_run(
    'We estimate EQUATION 1 using the first-difference estimator and present our regression '
    'results in TABLE 1.'
)
add_fn(p_res1,
       'The unabridged table with all coefficient on the controls along with a naive OLS regression '
       'without any control variables can be found in TABLE OA5 in the Appendix.')
p_res1.add_run(
    ' Depending on the specification, settlements where attacks took place saw an approximately '
    '9-14 percentage points larger increase in the vote share of Jobbik relative to the control '
    'group. The average baseline change in the control group is 15-21 percentage points. Thus, '
    'the estimated local impact of the attacks are substantial: the increase in Jobbik support is '
    '53% to 70% larger in attacked villages as it is in their respective control groups. Additional '
    'analyses - reported in the Online Appendix show that the attacks did not affect turnout; '
    'rather, they decreased the combined support of mainstream parties by an equal amount.'
)

p_res2 = doc.add_paragraph()
p_res2.add_run(
    'Using data from later elections we also test the persistence of our results. We regress the '
    'vote share of extreme right support across all elections within a ±10-year time window, '
    'relative to the attacks on time dummies, their interactions with a dummy indicating the attacks, '
    'and a set of settlement fixed effects. The coefficients of the interactions capture how the '
    'difference in Jobbik support between treated and control settlements evolved over time. We plot '
    'these in FIGURE 1. In Panel A, the control group is the donor pool, in Panel B, the control '
    'groups are the locations of the planned attacks. The results suggest that the effect of the '
    'attacks on the surge of Jobbik’s popular support was sustained even in 2014, though it '
    'dissipated by 2018.'
)

p_res3 = doc.add_paragraph()
p_res3.add_run(
    'The same analysis also provides additional evidence that the parallel trend assumption '
    'underlying our empirical analysis holds: differences between attacked and control settlements '
    'stayed small and insignificant in the period leading to the treatment.'
)
add_fn(p_res3, 'The underlying regression results are in TABLE OA7 of the Online Appendix.')
p_res3.add_run(
    ' Our estimates based on the synthetic control method also allow the exploration of how '
    'treatment effects vary across individual attacks. In the Online Appendix, we show that the '
    'magnitude of individual effects is largely consistent with anecdotal evidence on each '
    'individual incident. Specifically, the estimated effects are substantially smaller in the '
    'cases where the attacks were unsuccessful from the terrorists’ perspective.'
)
add_fn(p_res3,
       'As an additional robustness check we also considered the sensitivity of our results to '
       'some unobserved confounder driving both the probability of attacks and changes in the '
       'support of Jobbik (Cinelli and Hazlett 2020). As explained in the Online Appendix '
       'ONLINE APPENDIX SECTION G, it appears unlikely that such a confounder would explain '
       'our results.')

p_res4 = doc.add_paragraph()
p_res4.add_run(
    'So far, our analysis assumed that the local impact of terror was completely concentrated in '
    'villages where the attacks took place. This assumption makes our estimates conservative in '
    'the sense that if the impact of the terror attacks spilled over to nearby villages, then the '
    'control groups used for our estimates are partially contaminated. To explore spatial spillovers '
    'directly, in FIGURE 2 we visualize the relationship between both pre- and post-treatment Jobbik '
    'support and the distance from villages where attacks took place (left panel) and where they '
    'were planned but not executed (right-panel). Exploring these four relationships – i.e. '
    'pre vs. post treatment and planned vs. successful attacks – not only allows us to quantify '
    'the degree of spillover but also to construct placebo tests.'
)

p_res5 = doc.add_paragraph()
p_res5.add_run(
    'The key finding is that in 2010, distance from realized attacks becomes an influential predictor '
    'of variation in Jobbik support. The slope of the line is -.11 (s.e.:0.01), meaning that as one '
    'gets closer to attacked villages, Jobbik support is expected to rise by one percentage point in '
    'every 9 kilometers. On the contrary, the distance from future planned and realized attacks does '
    'not explain any variation in Jobbik support in 2006 and proximity to planned attacks does not '
    'predict increased Jobbik support in 2010 either. Taken together, this demonstrates that '
    'terrorist attacks committed against the Roma had a positive impact on far-right support in '
    'Hungary that spilled over regionally.'
)

# --- Section: Conclusion ---
h_con = doc.add_paragraph('Conclusion')
h_con.style = doc.styles['Heading 1']

p_con1 = doc.add_paragraph()
p_con1.add_run(
    'The key contribution of our study is to bring additional evidence on the table in the scholarly '
    'debate about the political impact of right-wing terrorism. The lives lost in ethnic violence '
    'surely reflect underlying tensions in divided societies but it would be tempting to think that '
    'outbursts of violence against ethnic minorities limit the credibility of parties campaigning on '
    'ethno-nationalist platforms. Instead, using our unique research design we find that right-wing '
    'terrorism in fact can '
)
p_con1.add_run('strengthen').italic = True
p_con1.add_run(
    ' the local support for these parties. This results contradicts recent research that finds that '
    'challenge the importance of physical proximity to such incidents '
    '(Agerberg and Sohlberg 2021; Böhmelt, Bove, and Nussio 2020).'
)

p_con2 = doc.add_paragraph()
p_con2.add_run(
    'We believe that our – relatively understudied setting – deserves attention on its '
    'own right, especially given the prevalence of anti-Roma hate crimes in Eastern Europe – '
    'including high profile events such as the 2009 Vítkov arson attack in Czechia, or the 2011 '
    'anti-Roma protests in Bulgaria (Amnesty International 2010; BBC News 2011). However, several '
    'factors might limit the generalizability of our findings to other cases. First, prejudice '
    'against the Roma minority in Hungary is quite strong and deep-seated '
    '(Simonovits, Kezdi, and Kardos 2018) which might make it easier for the ethnic majority to '
    '“relativize the murders by juxtaposing them with criminal acts committed by Roma or by '
    'highlighting the putatively legitimate motivations of the perpetrators” '
    '(Szombati 2018, p. xiv). At the same time, compared to some other cases with a history of '
    'ethnic violence and weaker institutions (Hager, Krakowski, and Schaub 2019) our results '
    'might understate the overall political impact of right-wing terrorism.'
)

p_con3 = doc.add_paragraph()
p_con3.add_run(
    'An important limitation of our research design is that it only permits the identification of '
    'the '
)
p_con3.add_run('local').italic = True
p_con3.add_run(
    ' impact of right wing terrorism. As our identification strategy relies on the comparison of '
    'localities with and without an attack taking place, we can only speculate about the broader, '
    'nationwide impact of these events. Existing research highlights the central importance of media '
    'coverage in shaping how citizens make sense of terror attacks '
    '(Shoshani and Slone 2008; Jacobs and van Spanje 2023; Böhmelt, Bove, and Nussio 2020), '
    'including the decision of media outlets to classify an incident as an act of terror '
    '(Huff and Kertzer 2018). Thus, it is quite possible that our analysis focusing on local '
    'effects understates the total impact of the attacks, especially that the nationwide uptake '
    'in the support for Jobbik coincided with the events we study and a corresponding media '
    'attention to ethnic tensions.'
)

p_con4 = doc.add_paragraph()
p_con4.add_run(
    'These lingering questions about both the scope conditions and causal mechanisms explaining '
    'our findings call for future research exploring more systematically the role of antecedent '
    'conditions on the impact of right wing terror-attacks. As more and more case-based research '
    'surfaces meta-analyses could explore how the link between right-wing terrorism and far-right '
    'voting is shaped by inter-group relations, the strength of political institutions as well as '
    'both the supply of and demand for news covering these incidents. In addition, to distinguish '
    'between the possible causal mechanisms that mediate the impact of terror-attacks future '
    'research could combine observational and experimental evidence to assess the relative '
    'importance of temporal and spatial proximity to acts of terror vis a vis exposure to various '
    'narratives of these events shapes public attitudes.'
)

# --- References ---
doc.add_page_break()
h_ref = doc.add_paragraph('References')
h_ref.style = doc.styles['Heading 1']

REFS_LIST = [
    "Agerberg, Mattias, and Jacob Sohlberg. 2021. “Personal Proximity and Reactions to Terrorism.” Comparative Political Studies 54(14): 2512–2545.",
    "Amnesty International. 2010. European Union: The Situation of Roma, 2010. London: Amnesty International.",
    "Auger, Vincent A. 2020. “Right-Wing Terror.” Perspectives on Terrorism 14(3): 87–97.",
    "BBC News. 2011. “Bulgarian Rally Links Roma to Organised Crime.” October 1. https://www.bbc.com/news/world-europe-15140291.",
    "Berrebi, Claude, and Esteban F. Klor. 2008. “Are Voters Sensitive to Terrorism? Direct Evidence from the Israeli Electorate.” American Political Science Review 102(3): 279–301.",
    "Böhmelt, Tobias, Vincenzo Bove, and Enzo Nussio. 2020. “Can Terrorism Abroad Influence Migration Attitudes at Home?” American Journal of Political Science 64(3): 437–451.",
    "Central Statistical Office (Központi Statisztikai Hivatal). 2001. 2001 Census of Hungary. https://www.nepszamlalas2001.hu.",
    "Central Statistical Office (Központi Statisztikai Hivatal). 2018. TSTAR (KSH Terület Statisztika).",
    "Cinelli, Carlos, and Chad Hazlett. 2020. “Making Sense of Sensitivity: Extending Omitted Variable Bias.” Journal of the Royal Statistical Society Series B: Statistical Methodology 82(1): 39–67.",
    "Cossu, Elena, and Caterina Froio. 2025. “How Terror Attacks Shape Political Agendas on Multiculturalism in France.” Politics 13: 9743.",
    "Godefroidt, Amélie. 2023. “How Terrorism Does (and Does Not) Affect Citizens’ Political Attitudes: A Meta-Analysis.” American Journal of Political Science 67(1): 22–38.",
    "Gyöngyösi, Győző, and Emil Verner. 2020. “Financial Crisis, Creditor-Debtor Conflict, and Populism.” The Journal of Finance 77(4): 2471–2523.",
    "Hager, Anselm, Krzysztof Krakowski, and Max Schaub. 2019. “Ethnic Riots and Prosocial Behavior: Evidence from Kyrgyzstan.” American Political Science Review 113(4): 1029–1044.",
    "Huff, Connor, and Joshua D. Kertzer. 2018. “How the Public Defines Terrorism.” American Journal of Political Science 62(1): 55–71.",
    "Jacobs, Laura, and Joost van Spanje. 2021. “Not All Terror Is Alike: How Right-Wing Extremist and Islamist Terror Threat Affect Anti-Immigration Party Support.” International Journal of Public Opinion Research 33(4): 737–755.",
    "Jacobs, Laura, and Joost van Spanje. 2023. “Who’s Afraid of Terror News? The Interplay between News Consumption Patterns, Personal Experiences and Fear of Terrorism.” Mass Communication and Society 26(3): 486–508.",
    "Jakobsson, Niklas, and Svein Blom. 2014. “Did the 2011 Terror Attacks in Norway Change Citizens’ Attitudes toward Immigrants?” International Journal of Public Opinion Research 26(4): 475–486.",
    "Jászberényi, Sándor. 2016. “A vádlottak célja magánhadsereg létrehozása és polgárháború kirobbantása volt.” 24.hu, October 9.",
    "Karácsony, Gergely, and Dániel Róna. 2011. “The Secret of Jobbik. Reasons behind the Rise of the Hungarian Radical Right.” Journal of East European & Asian Studies 2(1): 61–92.",
    "Krause, Werner, and Miku Matsunaga. 2023. “Does Right-Wing Violence Affect Public Support for Radical Right Parties? Evidence from Germany.” Comparative Political Studies 56(14): 2269–2305.",
    "Legewie, Joscha. 2013. “Terrorist Events and Attitudes toward Immigrants: A Natural Experiment.” American Journal of Sociology 118(5): 1199–1245.",
    "Marques, J. M., V. Y. Yzerbyt, and J.-P. Leyens. 1988. “The ‘Black Sheep Effect’: Extremity of Judgments towards Ingroup Members as a Function of Group Identification.” European Journal of Social Psychology 18(1): 1–16.",
    "Miklósi, Gábor. 2011. “Egy megdöbbentő gyilkosságsorozat részletei.” Index.hu, March 25.",
    "Montalvo, Jose G. 2011. “Voting after the Bombings: A Natural Experiment on the Effect of Terrorist Attacks on Democratic Elections.” Review of Economics and Statistics 93(4): 1146–1154.",
    "MTI. 2010. “Romagyilkosságok – A sajtóban is elhíresült falvakra támadtak. – HÁTTÉR.” Magyar Távirati Iroda, August 10.",
    "MTI. 2011. “Romagyilkosságok – Ipolytarnócon akartak újabb támadást elkövetni.” Magyar Távirati Iroda, April 12.",
    "National Election Commission. 2024. Election Results Database (1990–2024). https://www.valasztas.hu/1990-2024_eredmenyek.",
    "Pickard, Harry, Georgios Efthyvoulou, Vincenzo Bove, et al. 2022. “What’s Left after Right-Wing Extremism? The Effects on Political Orientation.” European Journal of Political Research.",
    "Sabet, Navid, Marius Liebald, and Guido Friebel. 2025. “Terrorism and Voting: The Rise of Right-Wing Populism in Germany.” American Economic Journal: Economic Policy 17(3): 407–440.",
    "Shoshani, Anat, and Michelle Slone. 2008. “The Drama of Media Coverage of Terrorism: Emotional and Attitudinal Impact on the Audience.” Studies in Conflict & Terrorism 31(7): 627–640.",
    "Simonovits, Gábor, Gabor Kezdi, and Peter Kardos. 2018. “Seeing the World through the Other’s Eye: An Online Intervention Reducing Ethnic Prejudice.” American Political Science Review 112(1): 186–193.",
    "Szombati, Kristof. 2018. The Revolt of the Provinces: Anti-Gypsyism and Right-Wing Politics in Hungary. Oxford and New York: Berghahn Books.",
    "Tamás, Gáspár Miklós. 2019. “A cigányok tíz éve.” Mérce.hu, February 23.",
    "Völker, Teresa. 2024. “How Terrorist Attacks Distort Public Debates: A Comparative Study of Right-Wing and Islamist Extremism.” Journal of European Public Policy 31(11): 3487–3514.",
    "Walters, Joanna, and Alvin Chang. 2021. “Far-Right Terror Poses Bigger Threat to US than Islamist Extremism Post-9/11.” The Guardian, September 8.",
]

for ref in REFS_LIST:
    rp = doc.add_paragraph(ref)
    rp.paragraph_format.left_indent = Pt(18)
    rp.paragraph_format.first_line_indent = Pt(-18)

# ── Save initial docx to bytes ────────────────────────────────
buf = io.BytesIO()
doc.save(buf)
buf.seek(0)

# ── Inject footnotes.xml into the zip ─────────────────────────
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_footnotes_xml(footnotes_list):
    lines = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<w:footnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
        ' xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"'
        ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
        ' xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"'
        ' xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"'
        ' xmlns:o="urn:schemas-microsoft-com:office:office"'
        ' xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
        ' xmlns:v="urn:schemas-microsoft-com:vml"'
        ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'
        ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'
        ' xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
        ' xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"'
        ' xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"'
        ' xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"'
        ' xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"'
        ' xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"'
        ' xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"'
        ' xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"'
        ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
        ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"'
        ' mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14">',
        # separator footnote
        '  <w:footnote w:type="separator" w:id="-1">',
        '    <w:p><w:r><w:separator/></w:r></w:p>',
        '  </w:footnote>',
        # continuation separator
        '  <w:footnote w:type="continuationSeparator" w:id="0">',
        '    <w:p><w:r><w:continuationSeparator/></w:r></w:p>',
        '  </w:footnote>',
    ]

    import xml.sax.saxutils as saxutils

    for i, text in enumerate(footnotes_list, start=1):
        escaped = saxutils.escape(text)
        lines.append(f'  <w:footnote w:id="{i}">')
        lines.append(f'    <w:p>')
        lines.append(f'      <w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>')
        lines.append(f'      <w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr><w:footnoteRef/></w:r>')
        lines.append(f'      <w:r><w:t xml:space="preserve"> {escaped}</w:t></w:r>')
        lines.append(f'    </w:p>')
        lines.append(f'  </w:footnote>')

    lines.append('</w:footnotes>')
    return '\n'.join(lines).encode('utf-8')


footnotes_xml_bytes = make_footnotes_xml(footnotes_data)

# Rebuild the zip with footnotes added
out_buf = io.BytesIO()
with zipfile.ZipFile(buf, 'r') as zin:
    with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == 'word/document.xml':
                # Ensure footnotes relationship is referenced – python-docx may not add it
                # (nothing to patch here; relationship is in .rels)
                zout.writestr(item, data)

            elif item.filename == 'word/_rels/document.xml.rels':
                # Add footnotes relationship if missing
                txt = data.decode('utf-8')
                if 'footnotes' not in txt.lower():
                    txt = txt.replace(
                        '</Relationships>',
                        '<Relationship Id="rId99" '
                        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
                        'Target="footnotes.xml"/></Relationships>'
                    )
                zout.writestr(item, txt.encode('utf-8'))

            elif item.filename == '[Content_Types].xml':
                txt = data.decode('utf-8')
                if 'footnotes' not in txt.lower():
                    txt = txt.replace(
                        '</Types>',
                        '<Override PartName="/word/footnotes.xml" '
                        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                        '</Types>'
                    )
                zout.writestr(item, txt.encode('utf-8'))

            else:
                zout.writestr(item, data)

        # Add footnotes.xml
        zout.writestr('word/footnotes.xml', footnotes_xml_bytes)

out_buf.seek(0)
with open(OUTPUT, 'wb') as f:
    f.write(out_buf.read())

print(f"Written: {OUTPUT}")
print(f"Footnotes: {len(footnotes_data)}")
